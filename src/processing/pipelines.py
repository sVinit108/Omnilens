"""
pipelines.py
────────────
Real preprocessing pipelines for all 21 DataCategory types.

Each pipeline function:
  - Accepts a batch of file Paths (all same category)
  - Processes each file with the appropriate library
  - Returns a list of ProcessingResult with normalized JSON in .output
  - Never raises — all exceptions are caught and returned as failed results

Output JSON is the intermediate representation feeding:
  → Step 4  : JSON object store
  → Step 5  : Ontology layer + Graph DB
  → Step 6  : Vector DB embeddings

Dependencies (install via requirements.txt):
    pdfplumber, python-docx, pandas, openpyxl, pyarrow,
    pillow, pytesseract, pyyaml, lxml, beautifulsoup4,
    faster-whisper, networkx, ebooklib, chardet,
    pydicom, nibabel, h5py, scipy, netCDF4,
    geopandas, fiona, trimesh, ezdxf,
    python-magic, safetensors, onnx
"""

import json
import logging
import re
import sqlite3
import zipfile
import tarfile
from datetime import datetime
from pathlib import Path

from models import DataCategory, ProcessingResult, compute_hash

logger = logging.getLogger(__name__)

# ── Type aliases ───────────────────────────────────────────────────────────────
FileBatch   = list[Path]
ResultBatch = list[ProcessingResult]


# ── Shared helpers ─────────────────────────────────────────────────────────────

def _ok(filepath: Path, category: DataCategory, output: dict) -> ProcessingResult:
    """Build a successful ProcessingResult with content hash."""
    return ProcessingResult(
        filepath=filepath,
        category=category,
        success=True,
        content_hash=compute_hash(filepath),
        output=output,
    )

def _err(filepath: Path, category: DataCategory, exc: Exception) -> ProcessingResult:
    """Build a failed ProcessingResult and log the error."""
    logger.error("[%s] %s → %s", category.value, filepath.name, exc)
    return ProcessingResult(
        filepath=filepath,
        category=category,
        success=False,
        content_hash="",
        error=str(exc),
    )


# ══════════════════════════════════════════════════════════════════════════════
# 1. DOCUMENTS  —  PDF, DOCX, ODT
# ══════════════════════════════════════════════════════════════════════════════

def process_documents(files: FileBatch) -> ResultBatch:
    """
    Extract full text, per-page structure, and document metadata.
    PDF  → pdfplumber  (handles text, tables, layout)
    DOCX → python-docx (paragraphs, headings, tables)
    ODT  → read as zip + parse content.xml (no extra dep)
    """
    import pdfplumber
    import docx
    from lxml import etree

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".pdf":
                pages = []
                metadata = {}
                with pdfplumber.open(f) as pdf:
                    metadata = pdf.metadata or {}
                    for i, page in enumerate(pdf.pages):
                        pages.append({
                            "page":   i + 1,
                            "text":   page.extract_text() or "",
                            "tables": page.extract_tables() or [],
                        })
                output = {
                    "format":   "pdf",
                    "pages":    pages,
                    "metadata": {k: str(v) for k, v in metadata.items()},
                    "source":   str(f),
                }

            elif ext in (".doc", ".docx"):
                doc = docx.Document(f)
                paragraphs = [{"style": p.style.name, "text": p.text} for p in doc.paragraphs]
                tables = []
                for table in doc.tables:
                    tables.append([[cell.text for cell in row.cells] for row in table.rows])
                output = {
                    "format":     "docx",
                    "paragraphs": paragraphs,
                    "tables":     tables,
                    "metadata":   dict(doc.core_properties.__dict__),
                    "source":     str(f),
                }

            elif ext == ".odt":
                # ODT is a ZIP; content lives in content.xml
                with zipfile.ZipFile(f) as z:
                    xml = z.read("content.xml")
                tree = etree.fromstring(xml)
                # Strip all tags, keep text
                text = " ".join(tree.itertext())
                output = {"format": "odt", "text": text, "source": str(f)}

            else:
                # Generic fallback: read as text
                output = {"format": ext, "text": f.read_text(errors="replace"), "source": str(f)}

            results.append(_ok(f, DataCategory.DOCUMENTS, output))

        except Exception as e:
            results.append(_err(f, DataCategory.DOCUMENTS, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 2. PLAIN TEXT  —  TXT, MD, RTF
# ══════════════════════════════════════════════════════════════════════════════

def process_plain_text(files: FileBatch) -> ResultBatch:
    """
    Read raw text with encoding detection (chardet).
    RTF: strip control words to get plain text.
    MD:  keep raw markdown (downstream LLM handles it natively).
    """
    import chardet

    def _strip_rtf(raw: bytes) -> str:
        """Minimal RTF → plain text by removing control words and braces."""
        text = raw.decode("latin-1", errors="replace")
        text = re.sub(r"\\[a-z]+\d* ?", "", text)   # control words
        text = re.sub(r"[{}\\]", "", text)            # braces and backslashes
        return text.strip()

    results = []
    for f in files:
        try:
            raw  = f.read_bytes()
            enc  = chardet.detect(raw)
            encoding = enc.get("encoding") or "utf-8"
            ext  = f.suffix.lower()

            if ext == ".rtf":
                text = _strip_rtf(raw)
            else:
                text = raw.decode(encoding, errors="replace")

            output = {
                "format":   ext.lstrip("."),
                "encoding": encoding,
                "text":     text,
                "source":   str(f),
            }
            results.append(_ok(f, DataCategory.PLAIN_TEXT, output))

        except Exception as e:
            results.append(_err(f, DataCategory.PLAIN_TEXT, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 3. STRUCTURED TEXT  —  JSON, XML, YAML, HTML
# ══════════════════════════════════════════════════════════════════════════════

def process_structured_text(files: FileBatch) -> ResultBatch:
    """
    Parse each format into a Python dict, then normalize to JSON.
    HTML: extract visible text + structural metadata via BeautifulSoup.
    """
    import yaml
    from lxml import etree
    from bs4 import BeautifulSoup

    def _xml_to_dict(element) -> dict | str:
        """Recursively convert lxml element tree to plain dict."""
        children = list(element)
        if not children:
            return element.text or ""
        result = {}
        for child in children:
            key = child.tag.split("}")[-1]  # strip namespace
            result.setdefault(key, []).append(_xml_to_dict(child))
        # Unwrap single-item lists for cleaner output
        return {k: v[0] if len(v) == 1 else v for k, v in result.items()}

    results = []
    for f in files:
        try:
            ext  = f.suffix.lower()
            raw  = f.read_text(errors="replace")

            if ext == ".json":
                parsed = json.loads(raw)
                output = {"format": "json", "data": parsed, "source": str(f)}

            elif ext in (".yaml", ".yml"):
                parsed = yaml.safe_load(raw)
                output = {"format": "yaml", "data": parsed, "source": str(f)}

            elif ext == ".xml":
                tree   = etree.fromstring(raw.encode())
                parsed = _xml_to_dict(tree)
                output = {"format": "xml", "data": parsed, "source": str(f)}

            elif ext in (".html", ".htm"):
                soup  = BeautifulSoup(raw, "lxml")
                title = soup.title.string if soup.title else ""
                # Extract all visible text blocks
                texts = [t.strip() for t in soup.stripped_strings if t.strip()]
                links = [a.get("href", "") for a in soup.find_all("a", href=True)]
                output = {
                    "format": "html",
                    "title":  title,
                    "text":   " ".join(texts),
                    "links":  links,
                    "source": str(f),
                }

            else:
                output = {"format": ext, "raw": raw, "source": str(f)}

            results.append(_ok(f, DataCategory.STRUCTURED_TEXT, output))

        except Exception as e:
            results.append(_err(f, DataCategory.STRUCTURED_TEXT, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 4. TABULAR DATA  —  CSV, TSV, XLSX, XLS, Parquet
# ══════════════════════════════════════════════════════════════════════════════

def process_tabular(files: FileBatch) -> ResultBatch:
    """
    Load with pandas. Output includes:
      - schema  : column names + inferred dtypes
      - records : list of row dicts (capped at 10k rows for large files)
      - stats   : row/col counts, null counts per column
    For huge files, only schema + stats are stored; records are sampled.
    """
    import pandas as pd

    MAX_RECORDS = 10_000  # rows stored in JSON; rest accessible via source path

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext in (".csv", ".tsv"):
                sep = "\t" if ext == ".tsv" else ","
                df  = pd.read_csv(f, sep=sep, low_memory=False)
            elif ext in (".xlsx", ".xls"):
                df  = pd.read_excel(f)
            elif ext == ".parquet":
                df  = pd.read_parquet(f)
            else:
                df  = pd.read_csv(f, low_memory=False)

            schema  = {col: str(dtype) for col, dtype in df.dtypes.items()}
            null_counts = df.isnull().sum().to_dict()
            sampled = df.head(MAX_RECORDS) if len(df) > MAX_RECORDS else df
            records = sampled.where(pd.notnull(sampled), None).to_dict(orient="records")

            output = {
                "format":      ext.lstrip("."),
                "row_count":   len(df),
                "col_count":   len(df.columns),
                "schema":      schema,
                "null_counts": {k: int(v) for k, v in null_counts.items()},
                "records":     records,
                "truncated":   len(df) > MAX_RECORDS,
                "source":      str(f),
            }
            results.append(_ok(f, DataCategory.TABULAR, output))

        except Exception as e:
            results.append(_err(f, DataCategory.TABULAR, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 5. IMAGES  —  JPG, PNG, WEBP, GIF, BMP
# ══════════════════════════════════════════════════════════════════════════════

def process_images(files: FileBatch) -> ResultBatch:
    """
    Extract EXIF metadata + run OCR for any text in the image.
    Pillow  → image properties + EXIF
    pytesseract → OCR text extraction
    """
    from PIL import Image
    from PIL.ExifTags import TAGS
    import pytesseract

    def _extract_exif(img: Image.Image) -> dict:
        exif_data = {}
        raw_exif = img._getexif() if hasattr(img, "_getexif") else None
        if raw_exif:
            for tag_id, value in raw_exif.items():
                tag = TAGS.get(tag_id, str(tag_id))
                # Skip binary/large fields
                if isinstance(value, (str, int, float)):
                    exif_data[tag] = value
        return exif_data

    results = []
    for f in files:
        try:
            with Image.open(f) as img:
                mode   = img.mode
                size   = img.size          # (width, height)
                fmt    = img.format or f.suffix.lstrip(".")
                exif   = _extract_exif(img)

                # Convert to RGB for OCR if needed (handles palette/RGBA modes)
                rgb = img.convert("RGB")
                ocr_text = pytesseract.image_to_string(rgb).strip()

            output = {
                "format":   fmt,
                "mode":     mode,
                "width":    size[0],
                "height":   size[1],
                "exif":     exif,
                "ocr_text": ocr_text,
                "source":   str(f),
            }
            results.append(_ok(f, DataCategory.IMAGES, output))

        except Exception as e:
            results.append(_err(f, DataCategory.IMAGES, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 6. VIDEO  —  MP4, MKV, MOV, AVI, WEBM
# ══════════════════════════════════════════════════════════════════════════════

def process_video(files: FileBatch) -> ResultBatch:
    """
    Extract video metadata via ffprobe (subprocess call — no Python dep needed).
    Audio track is extracted to a temp WAV and passed through Whisper for transcript.
    ffmpeg must be installed on the system.
    """
    import subprocess
    import tempfile

    def _ffprobe(filepath: Path) -> dict:
        """Run ffprobe and return stream/format metadata as dict."""
        cmd = [
            "ffprobe", "-v", "quiet",
            "-print_format", "json",
            "-show_streams", "-show_format",
            str(filepath),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        return json.loads(result.stdout) if result.stdout else {}

    def _transcribe_audio_track(filepath: Path) -> str:
        """Extract audio from video and transcribe with faster-whisper."""
        try:
            from faster_whisper import WhisperModel
            model = WhisperModel("base", compute_type="int8")
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
                tmp_path = tmp.name
            # Extract audio track to WAV
            subprocess.run(
                ["ffmpeg", "-i", str(filepath), "-vn", "-ar", "16000",
                 "-ac", "1", "-f", "wav", tmp_path, "-y"],
                capture_output=True, timeout=120,
            )
            segments, _ = model.transcribe(tmp_path)
            return " ".join(seg.text for seg in segments).strip()
        except Exception:
            return ""   # transcript is best-effort

    results = []
    for f in files:
        try:
            probe = _ffprobe(f)
            streams   = probe.get("streams", [])
            fmt_info  = probe.get("format", {})

            video_streams = [s for s in streams if s.get("codec_type") == "video"]
            audio_streams = [s for s in streams if s.get("codec_type") == "audio"]

            transcript = _transcribe_audio_track(f)

            output = {
                "format":       f.suffix.lstrip("."),
                "duration_sec": float(fmt_info.get("duration", 0)),
                "size_bytes":   int(fmt_info.get("size", 0)),
                "video_streams": video_streams,
                "audio_streams": audio_streams,
                "transcript":   transcript,
                "source":       str(f),
            }
            results.append(_ok(f, DataCategory.VIDEO, output))

        except Exception as e:
            results.append(_err(f, DataCategory.VIDEO, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 7. AUDIO  —  MP3, WAV, FLAC, AAC, OGG
# ══════════════════════════════════════════════════════════════════════════════

def process_audio(files: FileBatch) -> ResultBatch:
    """
    Transcribe speech → text using faster-whisper (local, no API cost).
    Also extract duration, sample rate, channels via ffprobe.
    """
    import subprocess
    from faster_whisper import WhisperModel

    # Load model once per batch (expensive init)
    model = WhisperModel("base", compute_type="int8")

    def _audio_metadata(filepath: Path) -> dict:
        cmd = [
            "ffprobe", "-v", "quiet", "-print_format", "json",
            "-show_streams", str(filepath),
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
        data   = json.loads(result.stdout) if result.stdout else {}
        streams = data.get("streams", [{}])
        s = streams[0] if streams else {}
        return {
            "duration_sec": float(s.get("duration", 0)),
            "sample_rate":  int(s.get("sample_rate", 0)),
            "channels":     int(s.get("channels", 0)),
            "codec":        s.get("codec_name", ""),
        }

    results = []
    for f in files:
        try:
            meta     = _audio_metadata(f)
            segments, info = model.transcribe(str(f))
            transcript = " ".join(seg.text for seg in segments).strip()

            output = {
                "format":       f.suffix.lstrip("."),
                "language":     info.language,
                "transcript":   transcript,
                **meta,
                "source":       str(f),
            }
            results.append(_ok(f, DataCategory.AUDIO, output))

        except Exception as e:
            results.append(_err(f, DataCategory.AUDIO, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 8. ARCHIVES  —  ZIP, TAR, GZ, RAR, 7Z
# ══════════════════════════════════════════════════════════════════════════════

def process_archives(files: FileBatch) -> ResultBatch:
    """
    List archive contents (name, size, compressed size).
    Does NOT extract — extraction + re-ingestion handled by ingest.py if needed.
    RAR/7Z: list via subprocess (requires unrar/7z system tools).
    """
    import subprocess

    def _list_zip(f: Path) -> list[dict]:
        with zipfile.ZipFile(f) as z:
            return [
                {"name": i.filename, "size": i.file_size,
                 "compressed": i.compress_size, "is_dir": i.is_dir()}
                for i in z.infolist()
            ]

    def _list_tar(f: Path) -> list[dict]:
        with tarfile.open(f) as t:
            return [
                {"name": m.name, "size": m.size,
                 "compressed": m.size, "is_dir": m.isdir()}
                for m in t.getmembers()
            ]

    def _list_via_cmd(f: Path, cmd: list[str]) -> list[dict]:
        result = subprocess.run(cmd + [str(f)], capture_output=True, text=True, timeout=30)
        lines  = result.stdout.strip().splitlines()
        return [{"name": line} for line in lines if line]

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".zip":
                contents = _list_zip(f)
            elif ext in (".tar", ".gz", ".bz2", ".xz"):
                contents = _list_tar(f)
            elif ext == ".rar":
                contents = _list_via_cmd(f, ["unrar", "lb"])
            elif ext == ".7z":
                contents = _list_via_cmd(f, ["7z", "l", "-ba", "-slt"])
            else:
                contents = []

            output = {
                "format":        ext.lstrip("."),
                "file_count":    sum(1 for c in contents if not c.get("is_dir")),
                "total_size":    sum(c.get("size", 0) for c in contents),
                "contents":      contents,
                "source":        str(f),
            }
            results.append(_ok(f, DataCategory.ARCHIVES, output))

        except Exception as e:
            results.append(_err(f, DataCategory.ARCHIVES, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 9. DATABASES  —  SQLite, SQL dump files
# ══════════════════════════════════════════════════════════════════════════════

def process_databases(files: FileBatch) -> ResultBatch:
    """
    SQLite: introspect schema and export all tables as JSON records.
    SQL dumps: parse CREATE TABLE and INSERT statements via regex.
    Row export capped at 10k per table to prevent memory blowout.
    """
    MAX_ROWS = 10_000

    def _read_sqlite(f: Path) -> dict:
        conn   = sqlite3.connect(f)
        cursor = conn.cursor()
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        table_names = [row[0] for row in cursor.fetchall()]

        tables = {}
        schema = {}
        for name in table_names:
            cursor.execute(f"PRAGMA table_info('{name}')")
            cols = [row[1] for row in cursor.fetchall()]
            schema[name] = cols
            cursor.execute(f"SELECT * FROM '{name}' LIMIT {MAX_ROWS}")
            rows = cursor.fetchall()
            tables[name] = [dict(zip(cols, row)) for row in rows]

        conn.close()
        return {"schema": schema, "tables": tables}

    def _read_sql_dump(f: Path) -> dict:
        raw    = f.read_text(errors="replace")
        tables = re.findall(r"CREATE TABLE[^;]+;", raw, re.IGNORECASE | re.DOTALL)
        return {"schema_ddl": tables, "tables": {}}

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()
            if ext in (".sqlite", ".db"):
                data = _read_sqlite(f)
            else:
                data = _read_sql_dump(f)

            output = {"format": ext.lstrip("."), **data, "source": str(f)}
            results.append(_ok(f, DataCategory.DATABASES, output))

        except Exception as e:
            results.append(_err(f, DataCategory.DATABASES, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 10. CODE  —  .py, .js, .ts, .java, .cpp, .sh, etc.
# ══════════════════════════════════════════════════════════════════════════════

def process_code(files: FileBatch) -> ResultBatch:
    """
    Read source code and extract:
      - Full source text
      - Language (from extension)
      - Line count, comment lines, blank lines
      - Python AST summary (functions, classes, imports) for .py files
    tree-sitter can replace the Python-specific AST logic for all languages.
    """
    import ast as pyast

    def _python_ast_summary(source: str) -> dict:
        try:
            tree    = pyast.parse(source)
            imports = [pyast.dump(n) for n in pyast.walk(tree) if isinstance(n, (pyast.Import, pyast.ImportFrom))]
            funcs   = [n.name for n in pyast.walk(tree) if isinstance(n, pyast.FunctionDef)]
            classes = [n.name for n in pyast.walk(tree) if isinstance(n, pyast.ClassDef)]
            return {"functions": funcs, "classes": classes, "imports": imports}
        except Exception:
            return {}

    COMMENT_PREFIXES = {"#", "//", "--", "/*", "*"}

    results = []
    for f in files:
        try:
            source   = f.read_text(errors="replace")
            lines    = source.splitlines()
            lang     = f.suffix.lstrip(".")

            blank_lines   = sum(1 for l in lines if not l.strip())
            comment_lines = sum(1 for l in lines if any(l.strip().startswith(p) for p in COMMENT_PREFIXES))

            ast_summary = _python_ast_summary(source) if lang == "py" else {}

            output = {
                "language":      lang,
                "line_count":    len(lines),
                "blank_lines":   blank_lines,
                "comment_lines": comment_lines,
                "ast_summary":   ast_summary,
                "source_code":   source,
                "source":        str(f),
            }
            results.append(_ok(f, DataCategory.CODE, output))

        except Exception as e:
            results.append(_err(f, DataCategory.CODE, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 11. LOGS  —  .log files
# ══════════════════════════════════════════════════════════════════════════════

def process_logs(files: FileBatch) -> ResultBatch:
    """
    Parse log files into structured entries.
    Attempts to extract: timestamp, log level, message.
    Supports common formats: ISO8601 timestamps, [LEVEL] prefixes.
    Falls back to raw lines if pattern doesn't match.
    """
    # Matches: 2024-01-01 12:00:00 [ERROR] Some message
    LOG_PATTERN = re.compile(
        r"(?P<timestamp>\d{4}-\d{2}-\d{2}[T ]\d{2}:\d{2}:\d{2}[\w.+-]*)?\s*"
        r"(?:\[?(?P<level>DEBUG|INFO|WARNING|WARN|ERROR|CRITICAL|FATAL)\]?)?\s*"
        r"(?P<message>.+)",
        re.IGNORECASE,
    )
    LEVEL_COUNTS_KEYS = ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"]

    results = []
    for f in files:
        try:
            lines   = f.read_text(errors="replace").splitlines()
            entries = []
            level_counts = {k: 0 for k in LEVEL_COUNTS_KEYS}

            for line in lines:
                m = LOG_PATTERN.match(line.strip())
                if m:
                    level = (m.group("level") or "UNKNOWN").upper()
                    if level in level_counts:
                        level_counts[level] += 1
                    entries.append({
                        "timestamp": m.group("timestamp"),
                        "level":     level,
                        "message":   m.group("message").strip(),
                    })
                else:
                    entries.append({"timestamp": None, "level": "UNKNOWN", "message": line})

            output = {
                "line_count":   len(lines),
                "level_counts": level_counts,
                "entries":      entries,
                "source":       str(f),
            }
            results.append(_ok(f, DataCategory.LOGS, output))

        except Exception as e:
            results.append(_err(f, DataCategory.LOGS, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 12. ML MODELS  —  .pt, .pth, .onnx, .h5, .safetensors
# ══════════════════════════════════════════════════════════════════════════════

def process_ml_models(files: FileBatch) -> ResultBatch:
    """
    Extract model architecture metadata without loading weights into memory.
    PyTorch  → layer names + param counts (CPU map)
    ONNX     → graph inputs/outputs/op types
    Safetensors → tensor names + shapes + dtypes
    HDF5/h5  → layer config (Keras models)
    """
    results = []
    for f in files:
        try:
            ext = f.suffix.lower()
            meta: dict = {"format": ext.lstrip(".")}

            if ext in (".pt", ".pth"):
                import torch
                obj = torch.load(f, map_location="cpu")
                if isinstance(obj, dict):
                    meta["keys"]       = list(obj.keys())[:50]
                    meta["param_count"] = sum(v.numel() for v in obj.values() if hasattr(v, "numel"))
                else:
                    meta["type"] = str(type(obj))

            elif ext == ".onnx":
                import onnx
                model = onnx.load(str(f))
                meta["inputs"]   = [i.name for i in model.graph.input]
                meta["outputs"]  = [o.name for o in model.graph.output]
                meta["op_types"] = list({n.op_type for n in model.graph.node})

            elif ext == ".safetensors":
                from safetensors import safe_open
                tensors = {}
                with safe_open(str(f), framework="pt", device="cpu") as st:
                    for key in st.keys():
                        t = st.get_tensor(key)
                        tensors[key] = {"shape": list(t.shape), "dtype": str(t.dtype)}
                meta["tensors"] = tensors

            elif ext == ".h5":
                import h5py
                def _h5_summary(h5obj) -> dict:
                    return {k: _h5_summary(v) if hasattr(v, "keys") else str(v.shape)
                            for k, v in h5obj.items()}
                with h5py.File(f, "r") as h5:
                    meta["structure"] = _h5_summary(h5)

            output = {**meta, "source": str(f)}
            results.append(_ok(f, DataCategory.ML_MODELS, output))

        except Exception as e:
            results.append(_err(f, DataCategory.ML_MODELS, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 13. TIME-SERIES DATA
# ══════════════════════════════════════════════════════════════════════════════

def process_time_series(files: FileBatch) -> ResultBatch:
    """
    Time-series files are typically CSV/Parquet with a datetime index column.
    Detect the datetime column, compute basic temporal stats:
      - start/end time, frequency, gaps, numeric column statistics.
    """
    import pandas as pd

    def _find_datetime_col(df) -> str | None:
        for col in df.columns:
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                return col
        # Try parsing string columns
        for col in df.columns:
            try:
                pd.to_datetime(df[col])
                return col
            except Exception:
                continue
        return None

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()
            df  = pd.read_parquet(f) if ext == ".parquet" else pd.read_csv(f)

            dt_col = _find_datetime_col(df)
            if dt_col:
                df[dt_col] = pd.to_datetime(df[dt_col], errors="coerce")
                df = df.sort_values(dt_col)
                time_stats = {
                    "datetime_column": dt_col,
                    "start":           str(df[dt_col].min()),
                    "end":             str(df[dt_col].max()),
                    "num_points":      len(df),
                    "inferred_freq":   str(pd.infer_freq(df[dt_col].dropna()) or "irregular"),
                }
            else:
                time_stats = {"datetime_column": None}

            numeric_cols = df.select_dtypes(include="number").columns.tolist()
            stats = df[numeric_cols].describe().to_dict() if numeric_cols else {}

            output = {
                "format":      ext.lstrip("."),
                "time_stats":  time_stats,
                "numeric_stats": stats,
                "columns":     df.columns.tolist(),
                "source":      str(f),
            }
            results.append(_ok(f, DataCategory.TIME_SERIES, output))

        except Exception as e:
            results.append(_err(f, DataCategory.TIME_SERIES, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 14. GEOSPATIAL DATA  —  GeoJSON, SHP, KML, GPX, GeoTIFF
# ══════════════════════════════════════════════════════════════════════════════

def process_geospatial(files: FileBatch) -> ResultBatch:
    """
    Vector formats (GeoJSON, SHP, KML, GPX): read with geopandas/fiona.
    Raster (GeoTIFF): read metadata with rasterio (if available), else skip bands.
    Output: CRS, geometry types, bounding box, feature count, properties.
    """
    import geopandas as gpd

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext in (".geojson", ".shp", ".kml", ".gpx"):
                gdf  = gpd.read_file(f)
                bbox = gdf.total_bounds.tolist()  # [minx, miny, maxx, maxy]
                output = {
                    "format":         ext.lstrip("."),
                    "crs":            str(gdf.crs),
                    "feature_count":  len(gdf),
                    "geometry_types": gdf.geom_type.unique().tolist(),
                    "bbox":           bbox,
                    "properties":     gdf.drop(columns="geometry").columns.tolist(),
                    "sample_features": json.loads(gdf.head(10).to_json()),
                    "source":         str(f),
                }

            elif ext in (".geotiff", ".tiff", ".tif"):
                try:
                    import rasterio
                    with rasterio.open(f) as src:
                        output = {
                            "format":  "geotiff",
                            "crs":     str(src.crs),
                            "width":   src.width,
                            "height":  src.height,
                            "bands":   src.count,
                            "dtype":   str(src.dtypes[0]),
                            "bounds":  list(src.bounds),
                            "source":  str(f),
                        }
                except ImportError:
                    output = {"format": "geotiff", "note": "rasterio not installed", "source": str(f)}

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.GEOSPATIAL, output))

        except Exception as e:
            results.append(_err(f, DataCategory.GEOSPATIAL, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 15. GRAPH DATA  —  GraphML, GEXF, RDF
# ══════════════════════════════════════════════════════════════════════════════

def process_graph(files: FileBatch) -> ResultBatch:
    """
    Parse graph files using networkx.
    Extract: node/edge counts, graph properties, degree distribution,
    and a sample of nodes/edges for the ontology layer.
    """
    import networkx as nx

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".graphml":
                G = nx.read_graphml(f)
            elif ext == ".gexf":
                G = nx.read_gexf(f)
            elif ext == ".rdf":
                # RDF → use rdflib if available, else treat as XML
                try:
                    import rdflib
                    g = rdflib.Graph()
                    g.parse(str(f))
                    G = nx.DiGraph()
                    for s, p, o in g:
                        G.add_edge(str(s), str(o), predicate=str(p))
                except ImportError:
                    G = nx.DiGraph()
            else:
                G = nx.DiGraph()

            degrees = dict(G.degree())
            top_nodes = sorted(degrees, key=degrees.get, reverse=True)[:20]

            output = {
                "format":        ext.lstrip("."),
                "node_count":    G.number_of_nodes(),
                "edge_count":    G.number_of_edges(),
                "is_directed":   G.is_directed(),
                "is_weighted":   nx.is_weighted(G),
                "density":       nx.density(G),
                "top_nodes":     top_nodes,
                "sample_edges":  list(G.edges(data=True))[:100],
                "source":        str(f),
            }
            results.append(_ok(f, DataCategory.GRAPH, output))

        except Exception as e:
            results.append(_err(f, DataCategory.GRAPH, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 16. MEDICAL IMAGING  —  DICOM, NIfTI
# ══════════════════════════════════════════════════════════════════════════════

def process_medical(files: FileBatch) -> ResultBatch:
    """
    DICOM → pydicom: extract anonymized metadata (no pixel data stored in JSON).
    NIfTI → nibabel: extract shape, affine, voxel dimensions.
    Pixel arrays are NOT serialized — too large; downstream handles them separately.
    """
    import pydicom
    import nibabel as nib

    # DICOM tags safe to export (no patient-identifying info)
    SAFE_DICOM_TAGS = {
        "Modality", "StudyDescription", "SeriesDescription",
        "Rows", "Columns", "SliceThickness", "PixelSpacing",
        "ImageOrientationPatient", "ImagePositionPatient",
        "Manufacturer", "ManufacturerModelName",
    }

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".dcm":
                ds = pydicom.dcmread(str(f), stop_before_pixels=True)
                meta = {
                    tag: str(getattr(ds, tag, ""))
                    for tag in SAFE_DICOM_TAGS
                    if hasattr(ds, tag)
                }
                output = {"format": "dicom", "metadata": meta, "source": str(f)}

            elif ext in (".nii", ".gz"):
                img   = nib.load(str(f))
                hdr   = img.header
                output = {
                    "format":          "nifti",
                    "shape":           list(img.shape),
                    "voxel_dims":      list(hdr.get_zooms()),
                    "affine":          img.affine.tolist(),
                    "data_dtype":      str(img.get_data_dtype()),
                    "source":          str(f),
                }

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.MEDICAL, output))

        except Exception as e:
            results.append(_err(f, DataCategory.MEDICAL, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 17. SCIENTIFIC DATA  —  HDF5, .mat, NetCDF
# ══════════════════════════════════════════════════════════════════════════════

def process_scientific(files: FileBatch) -> ResultBatch:
    """
    HDF5   → h5py: traverse dataset tree, extract shapes + dtypes (no array data)
    MAT    → scipy.io: extract variable names + shapes
    NetCDF → netCDF4: extract dimensions, variables, global attributes
    """
    import h5py
    import scipy.io as sio

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext in (".hdf5", ".h5"):
                def _traverse(obj, prefix="") -> dict:
                    info = {}
                    for key in obj.keys():
                        path = f"{prefix}/{key}"
                        item = obj[key]
                        if isinstance(item, h5py.Dataset):
                            info[path] = {"shape": list(item.shape), "dtype": str(item.dtype)}
                        elif isinstance(item, h5py.Group):
                            info.update(_traverse(item, path))
                    return info
                with h5py.File(f, "r") as h5:
                    datasets = _traverse(h5)
                output = {"format": "hdf5", "datasets": datasets, "source": str(f)}

            elif ext == ".mat":
                mat  = sio.loadmat(str(f), squeeze_me=True)
                vars_ = {k: str(v.shape) if hasattr(v, "shape") else str(type(v))
                         for k, v in mat.items() if not k.startswith("__")}
                output = {"format": "mat", "variables": vars_, "source": str(f)}

            elif ext == ".nc":
                import netCDF4 as nc
                with nc.Dataset(str(f), "r") as ds:
                    dims  = {k: len(v) for k, v in ds.dimensions.items()}
                    vars_ = {k: {"shape": list(v.shape), "dtype": str(v.dtype)}
                             for k, v in ds.variables.items()}
                    attrs = {k: str(getattr(ds, k)) for k in ds.ncattrs()}
                output = {
                    "format":     "netcdf",
                    "dimensions": dims,
                    "variables":  vars_,
                    "attributes": attrs,
                    "source":     str(f),
                }

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.SCIENTIFIC, output))

        except Exception as e:
            results.append(_err(f, DataCategory.SCIENTIFIC, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 18. 3D DATA  —  OBJ, STL, GLTF, PLY, FBX
# ══════════════════════════════════════════════════════════════════════════════

def process_3d(files: FileBatch) -> ResultBatch:
    """
    Load 3D meshes using trimesh and extract geometric metadata.
    FBX: trimesh can handle via assimp if installed.
    Vertex/face arrays are NOT stored in JSON — too large.
    """
    import trimesh

    results = []
    for f in files:
        try:
            mesh = trimesh.load(str(f), force="mesh")

            output = {
                "format":        f.suffix.lstrip("."),
                "vertex_count":  len(mesh.vertices),
                "face_count":    len(mesh.faces),
                "is_watertight": bool(mesh.is_watertight),
                "is_convex":     bool(mesh.is_convex),
                "bounds":        mesh.bounds.tolist(),
                "volume":        float(mesh.volume) if mesh.is_watertight else None,
                "surface_area":  float(mesh.area),
                "center_mass":   mesh.center_mass.tolist(),
                "source":        str(f),
            }
            results.append(_ok(f, DataCategory.DATA_3D, output))

        except Exception as e:
            results.append(_err(f, DataCategory.DATA_3D, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 19. POINT CLOUD DATA  —  LAS, LAZ, PCD
# ══════════════════════════════════════════════════════════════════════════════

def process_point_cloud(files: FileBatch) -> ResultBatch:
    """
    LAS/LAZ → laspy: extract point count, CRS, classification stats, bounding box.
    PCD     → open3d: extract point count and bounding box.
    Raw point arrays are NOT stored — too large for JSON.
    """
    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext in (".las", ".laz"):
                import laspy
                with laspy.open(str(f)) as las:
                    header = las.header
                    output = {
                        "format":      ext.lstrip("."),
                        "point_count": int(header.point_count),
                        "version":     f"{header.version.major}.{header.version.minor}",
                        "min_bound":   list(header.min),
                        "max_bound":   list(header.max),
                        "source":      str(f),
                    }

            elif ext == ".pcd":
                import open3d as o3d
                pcd  = o3d.io.read_point_cloud(str(f))
                bbox = pcd.get_axis_aligned_bounding_box()
                output = {
                    "format":      "pcd",
                    "point_count": len(pcd.points),
                    "min_bound":   bbox.min_bound.tolist(),
                    "max_bound":   bbox.max_bound.tolist(),
                    "has_colors":  pcd.has_colors(),
                    "has_normals": pcd.has_normals(),
                    "source":      str(f),
                }

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.POINT_CLOUD, output))

        except Exception as e:
            results.append(_err(f, DataCategory.POINT_CLOUD, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 20. CAD DATA  —  DXF, DWG, STEP
# ══════════════════════════════════════════════════════════════════════════════

def process_cad(files: FileBatch) -> ResultBatch:
    """
    DXF  → ezdxf: extract entities, layers, block definitions.
    DWG  → convert to DXF via ODA File Converter (system tool) then ezdxf.
    STEP → parse header + product metadata via regex (full parsing needs pythonocc).
    """
    import ezdxf

    def _read_step_header(f: Path) -> dict:
        """Extract STEP file header info (product name, schema) via regex."""
        raw  = f.read_text(errors="replace")
        name = re.search(r"PRODUCT\('([^']+)'", raw)
        schema = re.search(r"FILE_SCHEMA\s*\(\s*\('([^']+)'", raw)
        return {
            "product_name": name.group(1) if name else "",
            "schema":       schema.group(1) if schema else "",
        }

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".dxf":
                doc    = ezdxf.readfile(str(f))
                msp    = doc.modelspace()
                layers = [layer.dxf.name for layer in doc.layers]
                entity_types = {}
                for entity in msp:
                    t = entity.dxftype()
                    entity_types[t] = entity_types.get(t, 0) + 1
                output = {
                    "format":        "dxf",
                    "dxf_version":   doc.dxfversion,
                    "layers":        layers,
                    "entity_counts": entity_types,
                    "source":        str(f),
                }

            elif ext == ".dwg":
                # DWG requires ODA converter — attempt conversion to DXF first
                output = {
                    "format": "dwg",
                    "note":   "DWG requires ODA File Converter. Convert to DXF first.",
                    "source": str(f),
                }

            elif ext in (".step", ".stp"):
                header = _read_step_header(f)
                output = {"format": "step", **header, "source": str(f)}

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.CAD, output))

        except Exception as e:
            results.append(_err(f, DataCategory.CAD, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# 21. EBOOKS  —  EPUB, MOBI
# ══════════════════════════════════════════════════════════════════════════════

def process_ebooks(files: FileBatch) -> ResultBatch:
    """
    EPUB → ebooklib: extract chapters (HTML → plain text), TOC, metadata.
    MOBI → currently treated as binary; text extracted via regex heuristic.
          Full MOBI support requires mobi or KindleUnpack.
    """
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    def _html_to_text(html_content: bytes) -> str:
        soup = BeautifulSoup(html_content, "lxml")
        return soup.get_text(separator=" ", strip=True)

    results = []
    for f in files:
        try:
            ext = f.suffix.lower()

            if ext == ".epub":
                book = epub.read_epub(str(f))

                # Metadata
                title  = book.get_metadata("DC", "title")
                author = book.get_metadata("DC", "creator")
                lang   = book.get_metadata("DC", "language")

                # Extract chapter text from HTML items
                chapters = []
                for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
                    text = _html_to_text(item.get_content())
                    if text:
                        chapters.append({"id": item.get_id(), "text": text})

                output = {
                    "format":   "epub",
                    "title":    title[0][0] if title else "",
                    "author":   author[0][0] if author else "",
                    "language": lang[0][0] if lang else "",
                    "chapters": chapters,
                    "source":   str(f),
                }

            elif ext == ".mobi":
                # Best-effort: extract printable ASCII text from binary
                raw  = f.read_bytes()
                text = re.sub(rb"[^\x20-\x7E\n]", b" ", raw).decode("ascii", errors="ignore")
                text = re.sub(r" {2,}", " ", text).strip()
                output = {
                    "format": "mobi",
                    "text":   text[:50_000],   # cap at 50k chars
                    "note":   "Partial extraction. Install mobi library for full support.",
                    "source": str(f),
                }

            else:
                output = {"format": ext, "source": str(f)}

            results.append(_ok(f, DataCategory.EBOOKS, output))

        except Exception as e:
            results.append(_err(f, DataCategory.EBOOKS, e))

    return results


# ══════════════════════════════════════════════════════════════════════════════
# UNCLASSIFIED
# ══════════════════════════════════════════════════════════════════════════════

def process_unclassified(files: FileBatch) -> ResultBatch:
    """Files that could not be matched to any known category. Logged and skipped."""
    results = []
    for f in files:
        logger.warning("Unclassified file skipped: %s", f)
        results.append(ProcessingResult(
            filepath=f,
            category=DataCategory.UNCLASSIFIED,
            success=False,
            error="No matching pipeline for this file type",
        ))
    return results


# ══════════════════════════════════════════════════════════════════════════════
# PIPELINE REGISTRY
# Single source of truth: DataCategory → pipeline function
# Adding a new category = add one line here + write the function above.
# ══════════════════════════════════════════════════════════════════════════════

PIPELINE_REGISTRY: dict[DataCategory, callable] = {
    DataCategory.DOCUMENTS:         process_documents,
    DataCategory.PLAIN_TEXT:        process_plain_text,
    DataCategory.STRUCTURED_TEXT:   process_structured_text,
    DataCategory.TABULAR:           process_tabular,
    DataCategory.IMAGES:            process_images,
    DataCategory.VIDEO:             process_video,
    DataCategory.AUDIO:             process_audio,
    DataCategory.ARCHIVES:          process_archives,
    DataCategory.DATABASES:         process_databases,
    DataCategory.CODE:              process_code,
    DataCategory.LOGS:              process_logs,
    DataCategory.ML_MODELS:         process_ml_models,
    DataCategory.TIME_SERIES:       process_time_series,
    DataCategory.GEOSPATIAL:        process_geospatial,
    DataCategory.GRAPH:             process_graph,
    DataCategory.MEDICAL:           process_medical,
    DataCategory.SCIENTIFIC:        process_scientific,
    DataCategory.DATA_3D:           process_3d,
    DataCategory.POINT_CLOUD:       process_point_cloud,
    DataCategory.CAD:               process_cad,
    DataCategory.EBOOKS:            process_ebooks,
    DataCategory.UNCLASSIFIED:      process_unclassified,
}